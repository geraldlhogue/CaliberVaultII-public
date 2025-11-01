import os
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from markdown2 import markdown
from docx import Document

class MarkdownHandler(FileSystemEventHandler):
    def __init__(self, input_dir, output_dir):
        self.input_dir = input_dir
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def on_created(self, event):
        if event.is_directory:
            return
        if event.src_path.endswith('.md'):
            self.convert_md_to_docx(event.src_path)

    def on_modified(self, event):
        if event.is_directory:
            return
        if event.src_path.endswith('.md'):
            self.convert_md_to_docx(event.src_path)

    def convert_md_to_docx(self, md_path):
        try:
            # Read Markdown file
            with open(md_path, 'r', encoding='utf-8') as md_file:
                md_content = md_file.read()

            # Convert Markdown to HTML
            html_content = markdown(md_content)

            # Create Word document
            doc = Document()
            doc.add_heading(os.path.basename(md_path).replace('.md', ''), 0)

            # Add HTML content as plain text (basic conversion)
            for line in html_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

            # Save to Word directory
            output_filename = os.path.basename(md_path).replace('.md', '.docx')
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            print(f"Converted {md_path} to {output_path}")
        except Exception as e:
            print(f"Error converting {md_path}: {e}")

def start_watching():
    # Use the current directory as the input directory
    input_dir = os.getcwd()
    # Set Word directory as a subdirectory
    output_dir = os.path.join(input_dir, "Word")
    event_handler = MarkdownHandler(input_dir, output_dir)
    observer = Observer()
    observer.schedule(event_handler, input_dir, recursive=False)
    observer.start()
    print(f"Watching directory: {input_dir}")
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

if __name__ == "__main__":
    start_watching()
